import os
import sys
import tempfile
import logging
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# Ensure the ASR directory is in path so asr_google can be imported
sys.path.insert(0, os.path.dirname(__file__))
load_dotenv(os.path.join(os.path.dirname(__file__), ".env"))

from asr_google import transcribe

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(message)s",
)
logger = logging.getLogger("asr-server")

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 200 * 1024 * 1024  # 200 MB
CORS(app, resources={r"/api/*": {"origins": "http://localhost:3000"}})

ALLOWED_EXTENSIONS = {"mp3", "wav", "m4a", "ogg", "flac", "webm"}

def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def generate_word_document(bengali_transcript: str, english_transcript: str) -> io.BytesIO:
    """Generate a Word document with bilingual transcripts."""
    doc = Document()
    
    # Add title
    title = doc.add_heading("Transcript", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add Bengali section
    doc.add_heading("Bengali Transcript", level=1)
    bengali_para = doc.add_paragraph(bengali_transcript)
    bengali_para.paragraph_format.line_spacing = 1.5
    
    # Add section break
    doc.add_paragraph()
    
    # Add English section
    doc.add_heading("English Transcript", level=1)
    english_para = doc.add_paragraph(english_transcript)
    english_para.paragraph_format.line_spacing = 1.5
    
    # Save to BytesIO
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


@app.route("/api/transcribe", methods=["POST"])
def transcribe_endpoint():
    logger.info("Incoming transcription request")

    if "file" not in request.files:
        logger.warning("Transcription request missing file field")
        return jsonify({"error": "No file provided"}), 400

    file = request.files["file"]
    if not file.filename:
        logger.warning("Transcription request contained empty filename")
        return jsonify({"error": "No file selected"}), 400

    if not allowed_file(file.filename):
        logger.warning("Rejected unsupported file type: %s", file.filename)
        return jsonify({"error": f"Unsupported file type. Allowed: {', '.join(ALLOWED_EXTENSIONS)}"}), 400

    suffix = "." + file.filename.rsplit(".", 1)[1].lower()
    tmp_path = None
    try:
        logger.info("Accepted file %s", file.filename)
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            file.save(tmp.name)
            tmp_path = tmp.name

        result = transcribe(tmp_path)
        logger.info("Transcription completed for %s", file.filename)
        return jsonify({
            "transcript": {
                "bengali": result["bengali"],
                "english": result["english"]
            }
        })

    except Exception as exc:
        logger.exception("Transcription failed for %s", file.filename)
        return jsonify({"error": str(exc)}), 500

    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


@app.route("/api/download-word", methods=["POST"])
def download_word():
    """Generate and return a Word document with transcripts."""
    try:
        logger.info("Incoming DOCX download request")
        data = request.get_json()
        bengali_transcript = data.get("bengali", "")
        english_transcript = data.get("english", "")
        
        if not bengali_transcript and not english_transcript:
            return jsonify({"error": "No transcript provided"}), 400
        
        doc_io = generate_word_document(bengali_transcript, english_transcript)
        
        return send_file(
            doc_io,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="transcript.docx"
        )
    except Exception as exc:
        logger.exception("Word download generation failed")
        return jsonify({"error": str(exc)}), 500


@app.route("/api/health", methods=["GET"])
def health():
    return jsonify({"status": "ok"})


if __name__ == "__main__":
    logger.info("Starting ASR server on http://0.0.0.0:5001")
    app.run(debug=True, port=5001, host="0.0.0.0", threaded=True)
